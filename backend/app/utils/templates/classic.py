"""经典专业模板 - 黑灰白 + 居中标题 + 下划线 H2

适用: 国企 / 事业单位 / 传统行业 / 金融
布局: 单栏
字体: STSong-Light
bullet: reportlab bulletText='-' + Helvetica 渲染 + leftIndent
"""

from io import BytesIO
from typing import List, Dict, Any
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.lib.enums import TA_LEFT, TA_CENTER
from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer,
    KeepTogether,
)
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont

from app.utils.templates.base import BaseTemplate, StyleOptions
from app.utils.templates._pdf_helpers import (
    register_cn_font,
    escape_text,
    render_runs_html,
)


CN_FONT = register_cn_font()


class ClassicTemplate(BaseTemplate):
    name = "classic"
    display_name = "经典专业"
    description = "黑灰白配色 + 居中标题 + 下划线 H2,适合国企/金融/传统行业"

    # ---------- PDF ----------
    def render_pdf(self, tokens: List[Dict[str, Any]], options: StyleOptions) -> bytes:
        base_fs = options.font_size
        lh = options.line_height
        gap = options.section_gap

        # 字号阶梯(base ± 比例)
        fs_h1 = base_fs * 1.95  # ~20pt
        fs_h2 = base_fs * 1.30  # ~13.5pt
        fs_h3 = base_fs * 1.10  # ~11.5pt
        fs_body = base_fs
        fs_small = base_fs * 0.95

        # 行距
        lead_h1 = fs_h1 * lh
        lead_h2 = fs_h2 * lh
        lead_h3 = fs_h3 * lh
        lead_body = fs_body * lh

        # 模块间距(默认 1.0 = 1.0x base, 2.0 = 2x)
        sb_h1 = 6 * gap
        sa_h1 = 4 * gap
        sb_h2 = 8 * gap
        sa_h2 = 3 * gap
        sb_h3 = 4 * gap
        sa_h3 = 2 * gap
        sa_body = 2 * gap

        # 样式
        title_style = ParagraphStyle(
            "TplClassicTitle",
            fontName=CN_FONT,
            fontSize=fs_h1,
            leading=lead_h1,
            spaceBefore=sb_h1,
            spaceAfter=sa_h1,
            alignment=TA_CENTER,
            textColor="#1a1a1a",
        )
        h2_style = ParagraphStyle(
            "TplClassicH2",
            fontName=CN_FONT,
            fontSize=fs_h2,
            leading=lead_h2,
            spaceBefore=sb_h2,
            spaceAfter=sa_h2,
            textColor="#1a1a1a",
            borderWidth=0,
            borderPadding=0,
        )
        h3_style = ParagraphStyle(
            "TplClassicH3",
            fontName=CN_FONT,
            fontSize=fs_h3,
            leading=lead_h3,
            spaceBefore=sb_h3,
            spaceAfter=sa_h3,
            textColor="#2a2a2a",
        )
        body_style = ParagraphStyle(
            "TplClassicBody",
            fontName=CN_FONT,
            fontSize=fs_body,
            leading=lead_body,
            spaceAfter=sa_body,
        )
        bullet_style = ParagraphStyle(
            "TplClassicBullet",
            parent=body_style,
            leftIndent=18,
            bulletIndent=4,
            spaceAfter=2,
            fontSize=fs_body,
            leading=lead_body,
        )
        bullet2_style = ParagraphStyle(
            "TplClassicBullet2",
            parent=bullet_style,
            leftIndent=34,
        )

        buf = BytesIO()
        doc = SimpleDocTemplate(
            buf,
            pagesize=A4,
            leftMargin=2 * cm,
            rightMargin=2 * cm,
            topMargin=1.8 * cm,
            bottomMargin=1.8 * cm,
            title="AI 改写简历 - 经典专业",
        )

        story = self._build_story(
            tokens,
            title_style,
            h2_style,
            h3_style,
            body_style,
            bullet_style,
            bullet2_style,
        )
        doc.build(story)
        return buf.getvalue()

    def _build_story(
        self,
        tokens,
        title_style,
        h2_style,
        h3_style,
        body_style,
        bullet_style,
        bullet2_style,
    ):
        story = []
        h1_count = 0
        i = 0
        n = len(tokens)
        while i < n:
            tok = tokens[i]
            t = tok["type"]
            if t == "h1":
                h1_count += 1
                if h1_count == 1:
                    story.append(Paragraph(escape_text(tok["content"]), title_style))
                else:
                    # 后续 H1 当作 H2 处理(避免多 H1 破坏布局)
                    story.append(Paragraph(escape_text(tok["content"]), h2_style))
                i += 1
            elif t == "h2":
                # H2 标题不强制 KeepTogether,让下面的每个 H3 项目独立 KeepTogether
                story.append(Paragraph(escape_text(tok["content"]), h2_style))
                i += 1
                # 消费 H2 下面的所有简单块,但 H3 起的内容单独 KeepTogether
                while i < n and tokens[i]["type"] not in ("h1", "h2"):
                    if tokens[i]["type"] == "h3":
                        h3_block = self._consume_block(
                            tokens,
                            i,
                            h2_style,
                            h3_style,
                            body_style,
                            bullet_style,
                            bullet2_style,
                        )
                        if h3_block["elements"]:
                            story.append(KeepTogether(h3_block["elements"]))
                        i = h3_block["next_i"]
                    else:
                        block = self._consume_block(
                            tokens,
                            i,
                            h2_style,
                            h3_style,
                            body_style,
                            bullet_style,
                            bullet2_style,
                        )
                        if block["elements"]:
                            story.extend(block["elements"])
                        i = block["next_i"]
            else:
                # 单段落/列表也走统一消费
                block = self._consume_block(
                    tokens,
                    i,
                    h2_style,
                    h3_style,
                    body_style,
                    bullet_style,
                    bullet2_style,
                )
                if block["elements"]:
                    story.extend(block["elements"])
                i = block["next_i"]
        return story

    def _consume_h2_block(
        self, tokens, start, h2_style, h3_style, body_style, bullet_style, bullet2_style
    ):
        """从 H2 开始,到下一个 H2 之前整段打包(防分页拆断项目)"""
        elements = []
        first = tokens[start]
        elements.append(Paragraph(escape_text(first["content"]), h2_style))
        i = start + 1
        n = len(tokens)
        while i < n and tokens[i]["type"] not in ("h1", "h2"):
            block = self._consume_block(
                tokens, i, h2_style, h3_style, body_style, bullet_style, bullet2_style
            )
            if block["elements"]:
                elements.extend(block["elements"])
            i = block["next_i"]
        return {"elements": elements, "next_i": i}

    def _consume_block(
        self, tokens, start, h2_style, h3_style, body_style, bullet_style, bullet2_style
    ):
        """消费一个 block(可能是 h3+p+list 组合)"""
        elements = []
        i = start
        n = len(tokens)
        while i < n and tokens[i]["type"] in ("h3", "p", "ul", "ol", "blockquote"):
            tok = tokens[i]
            t = tok["type"]
            if t == "h3":
                elements.append(Paragraph(escape_text(tok["content"]), h3_style))
            elif t == "p":
                inline = render_runs_html(tok["runs"])
                if inline.strip():
                    elements.append(Paragraph(inline, body_style))
            elif t == "ul":
                for item in tok["items"]:
                    inline = render_runs_html(item["runs"])
                    if not inline.strip():
                        continue
                    level = 0
                    if item["runs"] and "level" in item["runs"][0]:
                        level = item["runs"][0]["level"] or 0
                    style = bullet2_style if level > 0 else bullet_style
                    elements.append(Paragraph(inline, style, bulletText="-"))
            elif t == "ol":
                for idx, item in enumerate(tok["items"]):
                    inline = render_runs_html(item["runs"])
                    if not inline.strip():
                        continue
                    level = 0
                    if item["runs"] and "level" in item["runs"][0]:
                        level = item["runs"][0]["level"] or 0
                    style = bullet2_style if level > 0 else bullet_style
                    elements.append(Paragraph(inline, style, bulletText=f"{idx+1}."))
            elif t == "blockquote":
                elements.append(
                    Paragraph("│ " + escape_text(tok["content"]), body_style)
                )
            i += 1
        return {"elements": elements, "next_i": i}

    # ---------- DOCX ----------
    def render_docx(
        self, doc: Document, tokens: List[Dict[str, Any]], options: StyleOptions
    ) -> None:
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        def set_cn_font(run, size_pt=None, bold=None, color_hex=None):
            rPr = run._element.get_or_add_rPr()
            rFonts = rPr.find(qn("w:rFonts"))
            if rFonts is None:
                rFonts = OxmlElement("w:rFonts")
                rPr.append(rFonts)
            rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
            rFonts.set(qn("w:ascii"), "Calibri")
            rFonts.set(qn("w:hAnsi"), "Calibri")
            if size_pt is not None:
                run.font.size = Pt(size_pt)
            if bold is not None:
                run.bold = bold
            if color_hex:
                r, g, b = (
                    int(color_hex[1:3], 16),
                    int(color_hex[3:5], 16),
                    int(color_hex[5:7], 16),
                )
                run.font.color.rgb = RGBColor(r, g, b)

        base_fs = options.font_size
        # docx 字号转 pt(round 半磅)
        fs_h1 = round(base_fs * 1.95 * 2) / 2
        fs_h2 = round(base_fs * 1.30 * 2) / 2
        fs_h3 = round(base_fs * 1.10 * 2) / 2

        h1_count = 0
        for tok in tokens:
            t = tok["type"]
            if t == "h1":
                h1_count += 1
                p = doc.add_heading("", level=1)
                run = p.add_run(escape_text(tok["content"]))
                set_cn_font(run, size_pt=fs_h1, bold=True, color_hex="#1a1a1a")
                if h1_count == 1:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif t == "h2":
                p = doc.add_heading("", level=2)
                run = p.add_run(escape_text(tok["content"]))
                set_cn_font(run, size_pt=fs_h2, bold=True, color_hex="#1a1a1a")
            elif t == "h3":
                p = doc.add_heading("", level=3)
                run = p.add_run(escape_text(tok["content"]))
                set_cn_font(run, size_pt=fs_h3, bold=True, color_hex="#2a2a2a")
            elif t == "p":
                p = doc.add_paragraph()
                _add_runs(
                    p, tok["runs"], base_fs, color="#1a1a1a", set_cn_font=set_cn_font
                )
            elif t == "ul":
                for item in tok["items"]:
                    p = doc.add_paragraph(style="List Bullet")
                    _add_runs(
                        p,
                        item["runs"],
                        base_fs,
                        color="#1a1a1a",
                        set_cn_font=set_cn_font,
                    )
            elif t == "ol":
                for item in tok["items"]:
                    p = doc.add_paragraph(style="List Number")
                    _add_runs(
                        p,
                        item["runs"],
                        base_fs,
                        color="#1a1a1a",
                        set_cn_font=set_cn_font,
                    )


def _add_runs(paragraph, runs, base_fs, color, set_cn_font):
    """把 runs 列表(每个含 text/bold/code 等)写入 paragraph"""
    from app.utils.templates._pdf_helpers import render_runs_html

    # 直接用 render_runs_html 但去掉外层 <para> 包裹
    # 简化:逐 run 写入,保留 bold/code
    for r in runs:
        text = r.get("text", "")
        if not text:
            continue
        # bold 在 runs 里?
        is_bold = bool(r.get("bold"))
        run = paragraph.add_run(text)
        set_cn_font(run, size_pt=base_fs, bold=is_bold, color_hex=color)
        if r.get("code"):
            run.font.name = "Consolas"
            run.font.size = base_fs - 0.5
