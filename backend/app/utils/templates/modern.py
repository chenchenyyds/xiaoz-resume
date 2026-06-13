"""现代简约模板 - 蓝色主色 + 左侧色条 + 蓝色 bullet

适用: 互联网 / 科技 / 创业公司 / 投资机构
布局: 单栏 + 左侧 4pt 蓝色装饰条
字体: STSong-Light
bullet: reportlab bulletText='▸' 蓝色 + Helvetica 渲染
"""

from io import BytesIO
from typing import List, Dict, Any
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import cm
from reportlab.lib.enums import TA_LEFT
from reportlab.platypus import (
    BaseDocTemplate,
    PageTemplate,
    Frame,
    Paragraph,
    Spacer,
    KeepTogether,
)

from app.utils.templates.base import BaseTemplate, StyleOptions
from app.utils.templates._pdf_helpers import (
    register_cn_font,
    escape_text,
    render_runs_html,
)


CN_FONT = register_cn_font()


# 颜色主题
PRIMARY = "#2563eb"
DARK = "#0f172a"
GRAY = "#475569"


def _draw_left_bar(canvas_obj, doc):
    """左侧 4pt 蓝色装饰条(每页画)"""
    canvas_obj.saveState()
    canvas_obj.setFillColor(PRIMARY)
    canvas_obj.rect(0, 0, 4, doc.pagesize[1], fill=1, stroke=0)
    canvas_obj.restoreState()


class ModernTemplate(BaseTemplate):
    name = "modern"
    display_name = "现代简约"
    description = "蓝色主色 + 左侧色条 + 蓝色 bullet,适合互联网/科技/创业公司"

    def render_pdf(self, tokens: List[Dict[str, Any]], options: StyleOptions) -> bytes:
        base_fs = options.font_size
        lh = options.line_height
        gap = options.section_gap

        fs_h1 = base_fs * 1.85
        fs_h2 = base_fs * 1.25
        fs_h3 = base_fs * 1.05
        fs_body = base_fs

        title_style = ParagraphStyle(
            "TplModernTitle",
            fontName=CN_FONT,
            fontSize=fs_h1,
            leading=fs_h1 * lh,
            spaceBefore=4 * gap,
            spaceAfter=2 * gap,
            alignment=TA_LEFT,
            textColor=DARK,
        )
        h2_style = ParagraphStyle(
            "TplModernH2",
            fontName=CN_FONT,
            fontSize=fs_h2,
            leading=fs_h2 * lh,
            spaceBefore=8 * gap,
            spaceAfter=3 * gap,
            textColor=PRIMARY,
        )
        h3_style = ParagraphStyle(
            "TplModernH3",
            fontName=CN_FONT,
            fontSize=fs_h3,
            leading=fs_h3 * lh,
            spaceBefore=4 * gap,
            spaceAfter=2 * gap,
            textColor=DARK,
        )
        body_style = ParagraphStyle(
            "TplModernBody",
            fontName=CN_FONT,
            fontSize=fs_body,
            leading=fs_body * lh,
            spaceAfter=2 * gap,
            textColor=GRAY,
        )
        bullet_style = ParagraphStyle(
            "TplModernBullet",
            parent=body_style,
            leftIndent=18,
            bulletIndent=4,
            spaceAfter=2,
            textColor=GRAY,
        )
        bullet2_style = ParagraphStyle(
            "TplModernBullet2",
            parent=bullet_style,
            leftIndent=34,
        )

        buf = BytesIO()
        doc = BaseDocTemplate(
            buf,
            pagesize=A4,
            leftMargin=2 * cm,
            rightMargin=2 * cm,
            topMargin=1.8 * cm,
            bottomMargin=1.8 * cm,
            title="AI 改写简历 - 现代简约",
        )
        frame = Frame(
            doc.leftMargin,
            doc.bottomMargin,
            doc.width,
            doc.height,
            id="normal",
        )
        template = PageTemplate(id="main", frames=frame, onPage=_draw_left_bar)
        doc.addPageTemplates([template])

        story = self._build_story(
            tokens,
            title_style,
            h2_style,
            h3_style,
            body_style,
            bullet_style,
            bullet2_style,
        )
        doc.build(story)
        return buf.getvalue()

    def _build_story(
        self,
        tokens,
        title_style,
        h2_style,
        h3_style,
        body_style,
        bullet_style,
        bullet2_style,
    ):
        story = []
        h1_count = 0
        i = 0
        n = len(tokens)
        while i < n:
            tok = tokens[i]
            t = tok["type"]
            if t == "h1":
                h1_count += 1
                if h1_count == 1:
                    story.append(Paragraph(escape_text(tok["content"]), title_style))
                else:
                    story.append(Paragraph(escape_text(tok["content"]), h2_style))
                i += 1
            elif t == "h2":
                story.append(Paragraph(escape_text(tok["content"]), h2_style))
                i += 1
                while i < n and tokens[i]["type"] not in ("h1", "h2"):
                    if tokens[i]["type"] == "h3":
                        h3_block = self._consume_block(
                            tokens,
                            i,
                            h2_style,
                            h3_style,
                            body_style,
                            bullet_style,
                            bullet2_style,
                        )
                        if h3_block["elements"]:
                            story.append(KeepTogether(h3_block["elements"]))
                        i = h3_block["next_i"]
                    else:
                        block = self._consume_block(
                            tokens,
                            i,
                            h2_style,
                            h3_style,
                            body_style,
                            bullet_style,
                            bullet2_style,
                        )
                        if block["elements"]:
                            story.extend(block["elements"])
                        i = block["next_i"]
            else:
                block = self._consume_block(
                    tokens,
                    i,
                    h2_style,
                    h3_style,
                    body_style,
                    bullet_style,
                    bullet2_style,
                )
                if block["elements"]:
                    story.extend(block["elements"])
                i = block["next_i"]
        return story

    def _consume_h2_block(
        self, tokens, start, h2_style, h3_style, body_style, bullet_style, bullet2_style
    ):
        elements = []
        first = tokens[start]
        elements.append(Paragraph(escape_text(first["content"]), h2_style))
        i = start + 1
        n = len(tokens)
        while i < n and tokens[i]["type"] not in ("h1", "h2"):
            block = self._consume_block(
                tokens, i, h2_style, h3_style, body_style, bullet_style, bullet2_style
            )
            if block["elements"]:
                elements.extend(block["elements"])
            i = block["next_i"]
        return {"elements": elements, "next_i": i}

    def _consume_block(
        self, tokens, start, h2_style, h3_style, body_style, bullet_style, bullet2_style
    ):
        elements = []
        i = start
        n = len(tokens)
        while i < n and tokens[i]["type"] in ("h3", "p", "ul", "ol", "blockquote"):
            tok = tokens[i]
            t = tok["type"]
            if t == "h3":
                elements.append(Paragraph(escape_text(tok["content"]), h3_style))
            elif t == "p":
                inline = render_runs_html(tok["runs"])
                if inline.strip():
                    elements.append(Paragraph(inline, body_style))
            elif t == "ul":
                for item in tok["items"]:
                    inline = render_runs_html(item["runs"])
                    if not inline.strip():
                        continue
                    level = 0
                    if item["runs"] and "level" in item["runs"][0]:
                        level = item["runs"][0]["level"] or 0
                    style = bullet2_style if level > 0 else bullet_style
                    elements.append(Paragraph(inline, style, bulletText="\u25b8"))
            elif t == "ol":
                for idx, item in enumerate(tok["items"]):
                    inline = render_runs_html(item["runs"])
                    if not inline.strip():
                        continue
                    level = 0
                    if item["runs"] and "level" in item["runs"][0]:
                        level = item["runs"][0]["level"] or 0
                    style = bullet2_style if level > 0 else bullet_style
                    elements.append(Paragraph(inline, style, bulletText=f"{idx+1}."))
            elif t == "blockquote":
                elements.append(
                    Paragraph("│ " + escape_text(tok["content"]), body_style)
                )
            i += 1
        return {"elements": elements, "next_i": i}

    # ---------- DOCX ----------
    def render_docx(
        self, doc: Document, tokens: List[Dict[str, Any]], options: StyleOptions
    ) -> None:
        from docx.shared import Pt, RGBColor
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        def set_cn_font(run, size_pt=None, bold=None, color_hex=None):
            rPr = run._element.get_or_add_rPr()
            rFonts = rPr.find(qn("w:rFonts"))
            if rFonts is None:
                rFonts = OxmlElement("w:rFonts")
                rPr.append(rFonts)
            rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
            rFonts.set(qn("w:ascii"), "Calibri")
            rFonts.set(qn("w:hAnsi"), "Calibri")
            if size_pt is not None:
                run.font.size = Pt(size_pt)
            if bold is not None:
                run.bold = bold
            if color_hex:
                r, g, b = (
                    int(color_hex[1:3], 16),
                    int(color_hex[3:5], 16),
                    int(color_hex[5:7], 16),
                )
                run.font.color.rgb = RGBColor(r, g, b)

        base_fs = options.font_size
        fs_h1 = round(base_fs * 1.85 * 2) / 2
        fs_h2 = round(base_fs * 1.25 * 2) / 2
        fs_h3 = round(base_fs * 1.05 * 2) / 2

        h1_count = 0
        for tok in tokens:
            t = tok["type"]
            if t == "h1":
                h1_count += 1
                p = doc.add_heading("", level=1)
                run = p.add_run(escape_text(tok["content"]))
                set_cn_font(run, size_pt=fs_h1, bold=True, color_hex=DARK)
            elif t == "h2":
                p = doc.add_heading("", level=2)
                run = p.add_run(escape_text(tok["content"]))
                set_cn_font(run, size_pt=fs_h2, bold=True, color_hex=PRIMARY)
            elif t == "h3":
                p = doc.add_heading("", level=3)
                run = p.add_run(escape_text(tok["content"]))
                set_cn_font(run, size_pt=fs_h3, bold=True, color_hex=DARK)
            elif t == "p":
                p = doc.add_paragraph()
                for r in tok["runs"]:
                    text = r.get("text", "")
                    if not text:
                        continue
                    run = p.add_run(text)
                    set_cn_font(run, size_pt=base_fs, color_hex=GRAY)
            elif t == "ul":
                for item in tok["items"]:
                    p = doc.add_paragraph(style="List Bullet")
                    for r in item["runs"]:
                        text = r.get("text", "")
                        if not text:
                            continue
                        run = p.add_run(text)
                        set_cn_font(run, size_pt=base_fs, color_hex=GRAY)
            elif t == "ol":
                for item in tok["items"]:
                    p = doc.add_paragraph(style="List Number")
                    for r in item["runs"]:
                        text = r.get("text", "")
                        if not text:
                            continue
                        run = p.add_run(text)
                        set_cn_font(run, size_pt=base_fs, color_hex=GRAY)
