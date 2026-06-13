"""双栏时间线模板 - 左侧 30% 个人信息/技能 + 右侧 70% 工作/项目

适用: 设计 / 产品 / 外企 / 简历需要强烈视觉冲击
布局: 双栏(左深色 + 右白底) — 用两个 Frame 实现,避免 Table split 死循环
字体: STSong-Light
bullet: reportlab bulletText='●' + Helvetica 渲染
"""

from io import BytesIO
from typing import List, Dict, Any
from docx import Document
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import cm
from reportlab.lib.enums import TA_LEFT
from reportlab.platypus import (
    BaseDocTemplate,
    PageTemplate,
    Frame,
    Paragraph,
    Spacer,
    KeepTogether,
)

from app.utils.templates.base import BaseTemplate, StyleOptions
from app.utils.templates._pdf_helpers import (
    register_cn_font,
    escape_text,
    render_runs_html,
)


CN_FONT = register_cn_font()


# 颜色
DARK_BG = "#1e293b"
DARK_TEXT = "#0f172a"
ACCENT = "#06b6d4"
WHITE = "#ffffff"
GRAY_LIGHT = "#cbd5e1"


def _draw_bg(canvas_obj, doc):
    """画左栏深色背景(每页,每页都画)"""
    canvas_obj.saveState()
    left_w = 6 * cm
    canvas_obj.setFillColor(DARK_BG)
    canvas_obj.rect(0, 0, left_w, doc.pagesize[1], fill=1, stroke=0)
    canvas_obj.restoreState()


class SidebarTemplate(BaseTemplate):
    name = "sidebar"
    display_name = "双栏时间线"
    description = "左侧深色栏(个人信息/技能) + 右侧白底(工作/项目),适合设计/产品/外企"

    LEFT_W = 6 * cm  # 左栏固定 6cm

    def render_pdf(self, tokens: List[Dict[str, Any]], options: StyleOptions) -> bytes:
        base_fs = options.font_size
        lh = options.line_height
        gap = options.section_gap

        fs_h1 = base_fs * 1.75
        fs_h2 = base_fs * 1.20
        fs_h3 = base_fs * 1.05
        fs_body = base_fs

        # 左栏样式(白字)
        left_title = ParagraphStyle(
            "TplSidebarLTitle",
            fontName=CN_FONT,
            fontSize=fs_h1,
            leading=fs_h1 * lh,
            spaceBefore=4 * gap,
            spaceAfter=2 * gap,
            textColor=WHITE,
        )
        left_h2 = ParagraphStyle(
            "TplSidebarLH2",
            fontName=CN_FONT,
            fontSize=fs_h2,
            leading=fs_h2 * lh,
            spaceBefore=6 * gap,
            spaceAfter=2 * gap,
            textColor=ACCENT,
        )
        left_body = ParagraphStyle(
            "TplSidebarLBody",
            fontName=CN_FONT,
            fontSize=fs_body * 0.95,
            leading=fs_body * 0.95 * lh,
            spaceAfter=1.5 * gap,
            textColor=GRAY_LIGHT,
        )
        left_bullet = ParagraphStyle(
            "TplSidebarLBullet",
            parent=left_body,
            leftIndent=12,
            bulletIndent=2,
            textColor=GRAY_LIGHT,
        )

        # 右栏样式(深色字)
        right_h2 = ParagraphStyle(
            "TplSidebarRH2",
            fontName=CN_FONT,
            fontSize=fs_h2,
            leading=fs_h2 * lh,
            spaceBefore=6 * gap,
            spaceAfter=2 * gap,
            textColor=DARK_TEXT,
        )
        right_h3 = ParagraphStyle(
            "TplSidebarRH3",
            fontName=CN_FONT,
            fontSize=fs_h3,
            leading=fs_h3 * lh,
            spaceBefore=4 * gap,
            spaceAfter=1.5 * gap,
            textColor=DARK_TEXT,
        )
        right_body = ParagraphStyle(
            "TplSidebarRBody",
            fontName=CN_FONT,
            fontSize=fs_body,
            leading=fs_body * lh,
            spaceAfter=2 * gap,
            textColor=DARK_TEXT,
        )
        right_bullet = ParagraphStyle(
            "TplSidebarRBullet",
            parent=right_body,
            leftIndent=16,
            bulletIndent=2,
        )

        # 解析 tokens:分流到左/右栏
        LEFT_KEYWORDS = {
            "专业技能",
            "教育经历",
            "个人信息",
            "联系方式",
            "技能",
            "证书",
            "语言能力",
            "兴趣爱好",
        }
        left_story = []
        right_story = []
        h1_count = 0
        i = 0
        n = len(tokens)
        while i < n:
            tok = tokens[i]
            t = tok["type"]
            if t == "h1":
                h1_count += 1
                if h1_count == 1:
                    left_story.append(
                        Paragraph(escape_text(tok["content"]), left_title)
                    )
                    if i + 1 < n and tokens[i + 1]["type"] == "p":
                        left_story.append(
                            Paragraph(
                                render_runs_html(tokens[i + 1]["runs"]), left_body
                            )
                        )
                        i += 2
                        continue
                else:
                    left_story.append(Paragraph(escape_text(tok["content"]), left_h2))
                i += 1
            elif t == "h2":
                is_left = any(kw in tok["content"] for kw in LEFT_KEYWORDS)
                if is_left:
                    left_story.append(Paragraph(escape_text(tok["content"]), left_h2))
                    i += 1
                    # 消费 h2 之后的所有简单块
                    while i < n and tokens[i]["type"] in (
                        "h3",
                        "p",
                        "ul",
                        "ol",
                        "blockquote",
                    ):
                        left_story.extend(
                            self._build_simple(tokens[i], left_body, left_bullet)
                        )
                        i += 1
                else:
                    right_story.append(Paragraph(escape_text(tok["content"]), right_h2))
                    i += 1
                    while i < n and tokens[i]["type"] in (
                        "h3",
                        "p",
                        "ul",
                        "ol",
                        "blockquote",
                    ):
                        right_story.extend(
                            self._build_simple(tokens[i], right_body, right_bullet)
                        )
                        i += 1
            else:
                # 单 p/ul 等放到左栏(一般是 H1 副信息未匹配的)
                left_story.extend(self._build_simple(tok, left_body, left_bullet))
                i += 1

        if not left_story:
            left_story = [Spacer(1, 0.1)]
        if not right_story:
            right_story = [Spacer(1, 0.1)]

        # 用 BaseDocTemplate + 双 Frame + onPage 画左栏背景
        buf = BytesIO()
        doc = BaseDocTemplate(
            buf,
            pagesize=A4,
            leftMargin=2 * cm,
            rightMargin=2 * cm,
            topMargin=1.5 * cm,
            bottomMargin=1.5 * cm,
            title="AI 改写简历 - 双栏时间线",
        )
        # 左 Frame:宽度 = LEFT_W, 在 page 左边
        left_frame = Frame(
            doc.leftMargin,
            doc.bottomMargin,
            self.LEFT_W - 0.5 * cm,  # 留点空隙给背景延伸
            doc.height,
            id="left",
            showBoundary=0,
            leftPadding=4,
            rightPadding=4,
        )
        # 右 Frame:宽度 = 总宽 - LEFT_W
        right_frame = Frame(
            doc.leftMargin + self.LEFT_W,
            doc.bottomMargin,
            doc.width - self.LEFT_W,
            doc.height,
            id="right",
            showBoundary=0,
            leftPadding=8,
            rightPadding=4,
        )
        template = PageTemplate(
            id="main", frames=[left_frame, right_frame], onPage=_draw_bg
        )
        doc.addPageTemplates([template])

        # FrameTemplate 要求单 story,且 frame 顺序消费
        # 但我们想 left_frame 渲染 left_story, right_frame 渲染 right_story
        # 用 NextPageTemplate / FrameBreak 来切换 frame
        from reportlab.platypus import FrameBreak

        story = left_story + [FrameBreak()] + right_story
        doc.build(story)
        return buf.getvalue()

    def _build_simple(self, tok, body_style, bullet_style):
        """单个 token 转成 flowable 列表"""
        t = tok["type"]
        out = []
        if t == "h3":
            out.append(
                Paragraph(
                    escape_text(tok["content"]),
                    ParagraphStyle(
                        "_sbh3",
                        parent=body_style,
                        fontSize=body_style.fontSize * 1.05,
                        textColor=body_style.textColor,
                        leading=body_style.fontSize * 1.05 * 1.4,
                    ),
                )
            )
        elif t == "p":
            inline = render_runs_html(tok["runs"])
            if inline.strip():
                out.append(Paragraph(inline, body_style))
        elif t == "ul":
            for item in tok["items"]:
                inline = render_runs_html(item["runs"])
                if not inline.strip():
                    continue
                out.append(Paragraph(inline, bullet_style, bulletText="\u25cf"))
        elif t == "ol":
            for idx, item in enumerate(tok["items"]):
                inline = render_runs_html(item["runs"])
                if not inline.strip():
                    continue
                out.append(Paragraph(inline, bullet_style, bulletText=f"{idx+1}."))
        elif t == "blockquote":
            out.append(Paragraph("│ " + escape_text(tok["content"]), body_style))
        return out

    # ---------- DOCX ----------
    def render_docx(
        self, doc: Document, tokens: List[Dict[str, Any]], options: StyleOptions
    ) -> None:
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        def set_cn_font(run, size_pt=None, bold=None, color_hex=None):
            rPr = run._element.get_or_add_rPr()
            rFonts = rPr.find(qn("w:rFonts"))
            if rFonts is None:
                rFonts = OxmlElement("w:rFonts")
                rPr.append(rFonts)
            rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
            rFonts.set(qn("w:ascii"), "Calibri")
            rFonts.set(qn("w:hAnsi"), "Calibri")
            if size_pt is not None:
                run.font.size = Pt(size_pt)
            if bold is not None:
                run.bold = bold
            if color_hex:
                r, g, b = (
                    int(color_hex[1:3], 16),
                    int(color_hex[3:5], 16),
                    int(color_hex[5:7], 16),
                )
                run.font.color.rgb = RGBColor(r, g, b)

        base_fs = options.font_size
        fs_h1 = round(base_fs * 1.75 * 2) / 2
        fs_h2 = round(base_fs * 1.20 * 2) / 2
        fs_h3 = round(base_fs * 1.05 * 2) / 2

        h1_count = 0
        for tok in tokens:
            t = tok["type"]
            if t == "h1":
                h1_count += 1
                p = doc.add_heading("", level=1)
                run = p.add_run(escape_text(tok["content"]))
                set_cn_font(run, size_pt=fs_h1, bold=True, color_hex=DARK_TEXT)
                if h1_count == 1:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif t == "h2":
                p = doc.add_heading("", level=2)
                run = p.add_run(escape_text(tok["content"]))
                set_cn_font(run, size_pt=fs_h2, bold=True, color_hex=ACCENT)
            elif t == "h3":
                p = doc.add_heading("", level=3)
                run = p.add_run(escape_text(tok["content"]))
                set_cn_font(run, size_pt=fs_h3, bold=True, color_hex=DARK_TEXT)
            elif t == "p":
                p = doc.add_paragraph()
                for r in tok["runs"]:
                    text = r.get("text", "")
                    if not text:
                        continue
                    run = p.add_run(text)
                    set_cn_font(run, size_pt=base_fs, color_hex=DARK_TEXT)
            elif t == "ul":
                for item in tok["items"]:
                    p = doc.add_paragraph(style="List Bullet")
                    for r in item["runs"]:
                        text = r.get("text", "")
                        if not text:
                            continue
                        run = p.add_run(text)
                        set_cn_font(run, size_pt=base_fs, color_hex=DARK_TEXT)
            elif t == "ol":
                for item in tok["items"]:
                    p = doc.add_paragraph(style="List Number")
                    for r in item["runs"]:
                        text = r.get("text", "")
                        if not text:
                            continue
                        run = p.add_run(text)
                        set_cn_font(run, size_pt=base_fs, color_hex=DARK_TEXT)
