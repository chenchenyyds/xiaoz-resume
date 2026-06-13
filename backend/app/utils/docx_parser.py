"""docx 解析/生成 - V1 支持 docx 输入 + 结构化 docx 输出

输出按 MD 解析后的 token 流渲染:
- H1/H2/H3 → 标题样式(微软雅黑,不同字号,加粗)
- 列表 → bullet/number list(支持嵌套)
- 段落内 **加粗** / `行内代码` → 用 runs 应用字体格式
- > 引用 → 斜体 + 灰色 + 左缩进
"""

from io import BytesIO
from typing import Optional
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from loguru import logger

from app.utils.md_parser import parse_md


def parse_docx(content: bytes) -> dict:
    """解析 docx 字节,返回 {text, paragraphs, sections}

    简化处理:把每段当一段,不做结构化(姓名/教育/工作分开)
    V1 改写是基于全文的,不做结构化也够用
    """
    try:
        doc = Document(BytesIO(content))
    except Exception as e:
        logger.exception("docx 解析失败")
        raise ValueError(f"docx 文件无法解析: {e}")

    paragraphs = []
    for p in doc.paragraphs:
        text = p.text.strip()
        if text:
            paragraphs.append(text)

    full_text = "\n".join(paragraphs)
    return {
        "text": full_text,
        "paragraphs": paragraphs,
        "sections": [],
        "char_count": len(full_text),
    }


def build_docx(md_text: str) -> bytes:
    """把 MD 转成结构化 docx 字节

    流程: MD -> parse_md tokens -> 按 token 类型分别用 docx styles 渲染
    """
    import time
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    t0 = time.time()
    doc = Document()

    # 1. 设置 Normal 样式默认字体(中文用微软雅黑,英文/数字用默认)
    _setup_default_style(doc)

    # 2. 解析 MD
    tokens = parse_md(md_text)
    logger.info(
        f"[docx_builder] parsed {len(tokens)} tokens from md len={len(md_text or '')}"
    )

    # 3. 渲染 tokens
    for tok in tokens:
        t = tok["type"]
        if t == "h1":
            # 用 docx 内置 Heading 1 样式(便于大纲/导航/TOC/无障碍)
            p = doc.add_heading(tok["content"], level=1)
            for run in p.runs:
                run.font.name = "Microsoft YaHei"
                rpr = run._element.get_or_add_rPr()
                rfonts = rpr.find(qn("w:rFonts"))
                if rfonts is None:
                    rfonts = OxmlElement("w:rFonts")
                    rpr.append(rfonts)
                rfonts.set(qn("w:eastAsia"), "Microsoft YaHei")
                rfonts.set(qn("w:ascii"), "Calibri")
                rfonts.set(qn("w:hAnsi"), "Calibri")
        elif t == "h2":
            p = doc.add_heading(tok["content"], level=2)
            for run in p.runs:
                run.font.name = "Microsoft YaHei"
                rpr = run._element.get_or_add_rPr()
                rfonts = rpr.find(qn("w:rFonts"))
                if rfonts is None:
                    rfonts = OxmlElement("w:rFonts")
                    rpr.append(rfonts)
                rfonts.set(qn("w:eastAsia"), "Microsoft YaHei")
                rfonts.set(qn("w:ascii"), "Calibri")
                rfonts.set(qn("w:hAnsi"), "Calibri")
        elif t == "h3":
            p = doc.add_heading(tok["content"], level=3)
            for run in p.runs:
                run.font.name = "Microsoft YaHei"
                rpr = run._element.get_or_add_rPr()
                rfonts = rpr.find(qn("w:rFonts"))
                if rfonts is None:
                    rfonts = OxmlElement("w:rFonts")
                    rpr.append(rfonts)
                rfonts.set(qn("w:eastAsia"), "Microsoft YaHei")
                rfonts.set(qn("w:ascii"), "Calibri")
                rfonts.set(qn("w:hAnsi"), "Calibri")
        elif t == "p":
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            _add_runs_to_paragraph(p, tok["runs"])
        elif t in ("ul", "ol"):
            for idx, item in enumerate(tok["items"]):
                p = _add_list_paragraph(doc, t, idx, item["runs"])
        elif t == "blockquote":
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(20)
            run = p.add_run(tok["content"])
            run.italic = True
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # 4. 页脚
    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    p.text = "本简历由 AI 辅助生成,求职者应自行核实真实性"

    buf = BytesIO()
    doc.save(buf)
    bytes_out = buf.getvalue()
    ms = int((time.time() - t0) * 1000)
    logger.info(
        f"[docx_builder] OK bytes={len(bytes_out)} tokens={len(tokens)} duration={ms}ms"
    )
    return bytes_out


def _setup_default_style(doc) -> None:
    """设置 Normal 样式:中文微软雅黑,英文 Calibri,10.5pt"""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)
    rpr = style.element.rPr
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    rfonts.set(qn("w:ascii"), "Calibri")
    rfonts.set(qn("w:hAnsi"), "Calibri")


def _add_runs_to_paragraph(p, runs: list) -> None:
    """把 runs 应用到段落上(支持 b/i/code 格式)"""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    for r in runs:
        text = r.get("text", "")
        if not text:
            continue
        run = p.add_run(text)
        fmt = r.get("fmt", "") or ""
        if "b" in fmt:
            run.bold = True
        if "i" in fmt:
            run.italic = True
        if "code" in fmt:
            # 行内代码:等宽字体 + 浅灰底
            run.font.name = "Consolas"
            rpr = run._element.get_or_add_rPr()
            rfonts = rpr.find(qn("w:rFonts"))
            if rfonts is None:
                rfonts = OxmlElement("w:rFonts")
                rpr.append(rfonts)
            rfonts.set(qn("w:ascii"), "Consolas")
            rfonts.set(qn("w:hAnsi"), "Consolas")
            rfonts.set(qn("w:eastAsia"), "Microsoft YaHei")
            # 加底纹
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), "F2F2F2")
            rpr.append(shd)
        if "link" in fmt:
            href = r.get("href", "")
            if href:
                run.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)
                run.font.underline = True


def _add_list_paragraph(doc, list_type: str, idx: int, runs: list):
    """加一个列表项,用 docx 内置 List Bullet / List Number 样式

    支持嵌套 level:用 numId 偏移实现,level=0 一级 bullet,level=1 二级 bullet
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    # 检测嵌套 level(从 runs[0] 的 level 字段读)
    level = 0
    if runs and "level" in runs[0]:
        level = runs[0]["level"] or 0

    # 选择样式
    style_name = f"List Bullet {level + 1}" if level > 0 else "List Bullet"
    if list_type == "ol":
        style_name = f"List Number {level + 1}" if level > 0 else "List Number"

    p = doc.add_paragraph(style=style_name)
    p.paragraph_format.space_after = Pt(2)

    # 内容用 _add_runs_to_paragraph 复用 runs(支持 b/i/code/link 格式)
    _add_runs_to_paragraph(p, runs)
    return p
